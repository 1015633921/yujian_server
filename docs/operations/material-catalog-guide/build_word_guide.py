from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
ASSET_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = ROOT / "docs" / "operations" / "宇涧运营后台-材料添加图文指南.docx"

# Preset: compact_reference_guide.
# Named brand overrides: preset blue headings are consistently replaced by
# Yustream green. Arial Unicode MS is used for both Latin and CJK glyphs so the
# same DOCX renders reliably in Word and the macOS LibreOffice QA pipeline.
# Body geometry, spacing, list indents, and table geometry remain preset values.
FONT_LATIN = "Arial Unicode MS"
FONT_CJK = "Arial Unicode MS"
INK = "1A211D"
MUTED = "5F6762"
GREEN = "214C3D"
GREEN_SOFT = "EAF1ED"
GOLD = "B28A4A"
GOLD_SOFT = "F7F0E4"
BLUE_SOFT = "E8EEF5"
BORDER = "D9DFDB"
WHITE = "FFFFFF"
CAUTION = "7A5A00"
RISK = "9B1C1C"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                edge.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, width in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"Table widths must total {CONTENT_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(
    run,
    size: float | None = None,
    color: str = INK,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size: float = 11, color: str = INK, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def set_paragraph_spacing(
    paragraph,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.25,
    keep_with_next: bool | None = None,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    if keep_with_next is not None:
        fmt.keep_with_next = keep_with_next


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, 18, 10, GREEN),
        "Heading 2": (13, 14, 7, GREEN),
        "Heading 3": (12, 10, 5, "365E50"),
    }
    for style_name, (size, before, after, color) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=9, color=MUTED)


def set_headers_and_footers(doc: Document) -> None:
    for section in doc.sections:
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        first_header = section.first_page_header
        p = first_header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line_spacing=1)
        run = p.add_run("宇涧 YUSTREAM")
        set_run_font(run, size=9, color=GOLD, bold=True)

        header = section.header
        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        set_table_geometry(table, [5400, 3960], indent_dxa=0)
        for cell in table.rows[0].cells:
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                set_cell_border(cell, **{edge: {"val": "nil"}})
            set_cell_margins(cell, {"top": 0, "bottom": 0, "start": 0, "end": 0})
        left = table.cell(0, 0).paragraphs[0]
        set_paragraph_spacing(left, after=0, line_spacing=1)
        set_run_font(left.add_run("宇涧 Yustream · 运营后台材料指南"), size=9, color=MUTED, bold=True)
        right = table.cell(0, 1).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(right, after=0, line_spacing=1)
        set_run_font(right.add_run("测试环境操作手册"), size=9, color=MUTED)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(p, after=0, line_spacing=1)
        set_run_font(p.add_run("Yustream Operations  ·  "), size=9, color=MUTED)
        add_field(p, "PAGE")


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_label(doc: Document, text: str, color: str = GOLD) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=8, line_spacing=1)
    run = p.add_run(text.upper())
    set_run_font(run, size=9.5, color=color, bold=True)


def add_title_page(doc: Document) -> None:
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=88, after=0, line_spacing=1)

    add_label(doc, "OPERATIONS PLAYBOOK  /  测试环境")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=10, line_spacing=1.05)
    run = p.add_run("宇涧运营后台\n材料添加图文指南")
    set_run_font(run, size=29, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=34, line_spacing=1.2)
    run = p.add_run("材料类型  ·  材料分类  ·  品种 / 款式  ·  SKU")
    set_run_font(run, size=13.5, color=MUTED)

    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340], indent_dxa=0)
    stages = [
        ("01", "材料类型"),
        ("02", "材料分类"),
        ("03", "品种 / 款式"),
        ("04", "SKU"),
    ]
    for idx, (num, label) in enumerate(stages):
        cell = table.cell(0, idx)
        set_cell_shading(cell, GREEN_SOFT if idx < 3 else GOLD_SOFT)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "8", "color": WHITE},
            bottom={"val": "single", "sz": "8", "color": WHITE},
            left={"val": "single", "sz": "8", "color": WHITE},
            right={"val": "single", "sz": "8", "color": WHITE},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4, line_spacing=1.1)
        set_run_font(p.add_run(f"{num}\n"), size=9, color=GOLD, bold=True)
        set_run_font(p.add_run(label), size=10.5, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=40, after=4, line_spacing=1.2)
    set_run_font(p.add_run("先建目录，再建 SKU"), size=16, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=46, line_spacing=1.35)
    set_run_font(
        p.add_run("目录描述材料“是什么”；尺寸、价格、库存与启停状态只在 SKU 中维护。"),
        size=11,
        color=MUTED,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=3, line_spacing=1.1)
    set_run_font(p.add_run("适用对象：运营、商品与素材维护人员"), size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line_spacing=1.1)
    set_run_font(p.add_run("版本 V1.0  ·  2026-07-15"), size=9.5, color=GOLD, bold=True)


def add_section_title(doc: Document, number: str, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph(style="Heading 1")
    set_paragraph_spacing(p, before=0, after=4, line_spacing=1.05, keep_with_next=True)
    set_run_font(p.add_run(f"{number}  "), size=11, color=GOLD, bold=True)
    set_run_font(p.add_run(title), size=18, color=GREEN, bold=True)
    if subtitle:
        sub = doc.add_paragraph()
        set_paragraph_spacing(sub, after=10, line_spacing=1.25, keep_with_next=True)
        set_run_font(sub.add_run(subtitle), size=10.5, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 2) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_font(p, size=13 if level == 2 else 12, color=GREEN if level == 2 else "365E50", bold=True)
    p.add_run(text) if not p.runs else None
    if p.runs:
        set_paragraph_font(p, size=13 if level == 2 else 12, color=GREEN if level == 2 else "365E50", bold=True)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, color: str = INK) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6, line_spacing=1.25)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), size=11, color=color, bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]), size=11, color=color)
    else:
        set_run_font(p.add_run(text), size=11, color=color)


def add_bullet(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)
    set_paragraph_spacing(p, after=4, line_spacing=1.25)
    set_run_font(p.add_run(text), size=10.5, color=INK)


def add_numbered(doc: Document, text: str, num_id: int) -> None:
    add_bullet(doc, text, num_id)


def create_numbering(doc: Document, fmt: str, text: str) -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract, default=-1) + 1
    existing_num = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.extend([tabs, ind])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_callout(doc: Document, title: str, text: str, tone: str = "info") -> None:
    fill = GREEN_SOFT if tone == "info" else GOLD_SOFT
    accent = GREEN if tone == "info" else CAUTION
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "18", "color": accent},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "nil"},
    )
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=3, after=2, line_spacing=1.2)
    set_run_font(p.add_run(f"{title}  "), size=10.5, color=accent, bold=True)
    set_run_font(p.add_run(text), size=10.5, color=INK)
    tail = doc.add_paragraph()
    set_paragraph_spacing(tail, after=2, line_spacing=1)


def add_picture(doc: Document, filename: str, caption: str, alt_text: str, width: float = 6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=3, line_spacing=1)
    run = p.add_run()
    shape = run.add_picture(str(ASSET_DIR / filename), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, after=8, line_spacing=1.1)
    set_run_font(cap.add_run(caption), size=9, color=MUTED, italic=True)


def add_text_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, BLUE_SOFT)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": BORDER},
            bottom={"val": "single", "sz": "6", "color": BORDER},
            left={"val": "single", "sz": "6", "color": BORDER},
            right={"val": "single", "sz": "6", "color": BORDER},
        )
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=2, after=2, line_spacing=1.1)
        set_run_font(p.add_run(header), size=9.5, color=GREEN, bold=True)

    for row_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, WHITE if row_idx % 2 == 0 else "FAFBFA")
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": BORDER},
                bottom={"val": "single", "sz": "4", "color": BORDER},
                left={"val": "single", "sz": "4", "color": BORDER},
                right={"val": "single", "sz": "4", "color": BORDER},
            )
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=1.5, after=1.5, line_spacing=1.12)
            set_run_font(p.add_run(text), size=9.25, color=INK)

    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=1, line_spacing=1)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_document() -> Document:
    doc = Document()
    style_document(doc)
    bullet_num_id = create_numbering(doc, "bullet", "•")
    category_num_id = create_numbering(doc, "decimal", "%1.")
    variety_num_id = create_numbering(doc, "decimal", "%1.")
    sku_num_id = create_numbering(doc, "decimal", "%1.")

    core = doc.core_properties
    core.title = "宇涧运营后台材料添加图文指南"
    core.subject = "材料类型、材料分类、品种/款式与 SKU 的运营操作指南"
    core.author = "宇涧 Yustream"
    core.keywords = "宇涧, 运营后台, 材料, SKU, 图文指南"
    core.comments = "适用于测试环境运营后台。"

    add_title_page(doc)

    add_page_break(doc)
    add_section_title(doc, "01", "先理解四级目录", "所有材料都按同一条路径建立，先判断层级，再开始录入。")
    add_picture(
        doc,
        "01-directory-flow.png",
        "图 1  材料从目录到 SKU 的标准录入顺序",
        "材料类型、材料分类、品种款式、SKU 四级录入流程示意图",
        width=6.3,
    )
    add_callout(doc, "核心规则", "目录描述材料是什么；尺寸、价格、库存和启停状态只填写在 SKU。")
    add_text_table(
        doc,
        ["材料类型", "材料分类", "品种 / 款式", "SKU 示例"],
        [
            ["珠子", "幽灵水晶", "绿幽灵", "绿幽灵 8mm AAA"],
            ["珠子", "幽灵水晶", "红幽灵", "红幽灵 10mm AA"],
            ["配饰", "幽灵随形", "绿幽灵随形", "绿幽灵随形 #001"],
            ["配饰", "隔珠", "圆饼隔珠", "圆饼隔珠 4×2mm"],
            ["配饰", "吊坠", "月亮吊坠", "月亮吊坠 12×18mm"],
        ],
        [1440, 1944, 2376, 3600],
    )

    add_page_break(doc)
    add_section_title(doc, "02", "创建材料类型", "只建立最上层的业务边界，目前主要复用“珠子”和“配饰”。")
    add_picture(
        doc,
        "02-create-type.png",
        "图 2  材料类型页面与新增类型入口",
        "运营后台材料类型页面，标注新增类型按钮和类型编码填写区域",
        width=6.2,
    )
    add_heading(doc, "填写规则")
    for item in [
        "已有“珠子”和“配饰”时直接复用，不要重复创建。",
        "类型名称可以调整；类型编码创建后不可修改。",
        "类型编码使用小写英文，例如 bead、accessory。",
        "“幽灵随形”“隔珠”“花托”都不是材料类型，应放在配饰下的分类。",
    ]:
        add_bullet(doc, item, bullet_num_id)
    add_callout(doc, "判断方法", "如果某个名称不能同时容纳多个分类，就不应把它建成材料类型。", tone="caution")

    add_page_break(doc)
    add_section_title(doc, "03", "创建材料分类", "分类描述同一材料类型下的业务大类，不包含颜色、尺寸、等级、价格或库存。")
    add_picture(
        doc,
        "03-create-category.png",
        "图 3  在指定材料类型下创建材料分类",
        "材料分类新增弹窗，标注所属类型、分类名称、排序和启用状态",
        width=6.3,
    )
    add_heading(doc, "操作顺序")
    for item in [
        "选择材料类型，例如“配饰”。",
        "填写分类名称，例如“幽灵随形”。",
        "设置排序与启用状态。",
        "点击“保存分类”。",
    ]:
        add_numbered(doc, item, category_num_id)
    add_callout(doc, "命名边界", "“幽灵水晶”适合作为珠子分类；“绿幽灵”属于其下的品种，不要建成同级分类。")

    add_page_break(doc)
    add_section_title(doc, "04", "创建品种 / 款式", "品种页面采用级联选择，必须先选类型，再选该类型下的分类。")
    add_picture(
        doc,
        "04-create-variety.png",
        "图 4  通过级联选项创建品种或款式",
        "品种款式新增弹窗，标注材料类型、材料分类和品种名称的级联关系",
        width=6.3,
    )
    add_heading(doc, "操作顺序")
    for item in [
        "先选择材料类型，例如“配饰”。",
        "所属材料分类只会显示该类型下的分类。",
        "选择“幽灵随形”，再填写“绿幽灵随形”。",
        "点击“保存品种 / 款式”。",
    ]:
        add_numbered(doc, item, variety_num_id)
    add_callout(doc, "级联保护", "切换材料类型后，原分类若不属于新类型会自动清空；保存前必须重新选择。", tone="caution")

    add_page_break(doc)
    add_section_title(doc, "05", "完善品种资料", "品种资料由同一品种下的 SKU 共用，负责工作台表现与推荐解释。")
    add_picture(
        doc,
        "05-complete-profile.png",
        "图 5  品种保存后进入“完善资料”",
        "品种列表与完善资料入口，展示视觉素材、工作台形制和安装方式等共享字段",
        width=6.3,
    )
    add_heading(doc, "需要维护")
    for item in [
        "视觉素材：透明背景主图和详情图，按素材管线处理后上传 COS。",
        "工作台形制：圆珠、随形、双尖、单尖、方糖、花托或吊坠。",
        "安装方式：按实物选择穿线、悬挂或其他安装方式。",
        "穿线轴角度：原图横向穿线填 0，竖向穿线填 90。",
        "推荐资料：元素倾向、适用愿景、材料角色和搭配规则。",
        "养护信息：易磕碰、避免暴晒、建议分开收纳等。",
    ]:
        add_bullet(doc, item, bullet_num_id)
    add_callout(doc, "不要混填", "品种资料不保存价格和库存；这些销售字段只属于具体 SKU。")

    add_page_break(doc)
    add_section_title(doc, "06", "创建具体 SKU", "每一个可售规格或独立库存单位都必须建立自己的 SKU。")
    add_picture(
        doc,
        "06-create-sku.png",
        "图 6  SKU 页面中的目录选择、销售字段与实物规格",
        "新增 SKU 页面，标注类型分类品种级联、展示名称、价格库存和非圆珠实物规格",
        width=5.9,
    )
    add_heading(doc, "录入顺序")
    for item in [
        "依次选择材料类型、材料分类和品种 / 款式。",
        "填写展示名称、单颗价格、重量和库存。",
        "填写珠径或外观最大尺寸；非圆珠继续填写工作台实物规格。",
        "资料未验收前保持停用；工作台显示正常后再启用。",
    ]:
        add_numbered(doc, item, sku_num_id)

    add_page_break(doc)
    add_section_title(doc, "07", "规格怎么填", "圆珠与非圆珠使用不同的测量逻辑，不要只用一个“珠径”字段代替所有尺寸。")
    add_heading(doc, "圆珠")
    add_text_table(
        doc,
        ["字段", "填写方式"],
        [
            ["珠径 / 外观最大尺寸", "填写实际珠径，例如 8.0mm。"],
            ["穿线方向占位", "可留空，系统沿用珠径。"],
            ["外观宽度 / 高度", "可留空，系统按圆珠处理。"],
            ["库存", "填写该规格实际可售颗数。"],
        ],
        [2700, 6660],
    )
    add_callout(doc, "拆分规则", "同一品种的 8mm、10mm、12mm 必须分别建立 SKU。")

    add_heading(doc, "随形、双尖、单尖与其他非圆珠")
    add_text_table(
        doc,
        ["字段", "填写方式"],
        [
            ["珠径 / 外观最大尺寸", "填写该颗材料的最大外观尺寸。"],
            ["穿线方向占位", "填写穿在线上实际占用的宽度，决定成串间距。"],
            ["外观宽度", "填写图片横向对应的真实宽度。"],
            ["外观高度", "填写图片纵向对应的真实高度。"],
            ["库存", "每颗单独测量、单独管理时填写 1。"],
        ],
        [2700, 6660],
    )
    add_callout(
        doc,
        "批次与单颗",
        "同一批次的平均尺寸和公差写入批次或采购备注；每颗独立 SKU 仍填写该颗实测值。",
        tone="caution",
    )

    add_page_break(doc)
    add_section_title(doc, "08", "配饰专项规则", "配饰的工作台真实性取决于实物规格、安装方式与透明素材三者一致。")
    add_text_table(
        doc,
        ["配饰类型", "重点填写", "验收重点"],
        [
            ["隔珠", "穿线方向占位、外观宽度、外观高度", "成串间距与实物厚度一致"],
            ["花托", "包裹方向、外观尺寸、适配珠径", "不能悬空或压入珠体"],
            ["吊坠", "悬挂安装方式、连接点、外观尺寸", "不能按普通圆珠贴在线圈上"],
            ["合金配饰", "透明主图、穿线轴角度、实物尺寸", "无白底、无比例失真"],
            ["幽灵随形 / 双尖 / 单尖", "单颗实测尺寸、库存 1、正确形制", "孔位与穿线方向准确"],
        ],
        [1700, 3500, 4160],
    )
    add_heading(doc, "推荐目录示例")
    add_text_table(
        doc,
        ["目录路径", "SKU 示例"],
        [
            ["配饰 > 幽灵随形 > 红幽灵随形", "红幽灵随形 #001"],
            ["配饰 > 幽灵随形 > 绿幽灵随形", "绿幽灵随形 #001"],
            ["配饰 > 幽灵双尖 > 红幽灵双尖", "红幽灵双尖 #001"],
            ["配饰 > 隔珠 > 圆饼隔珠", "圆饼隔珠 4×2mm"],
            ["配饰 > 花托 > 莲花花托", "莲花花托 8mm"],
            ["配饰 > 吊坠 > 月亮吊坠", "月亮吊坠 12×18mm"],
        ],
        [5400, 3960],
    )
    add_callout(doc, "素材要求", "新增珠子与配饰图片必须经过统一素材管线并上传 COS，不能直接使用临时抠图。", tone="caution")

    add_page_break(doc)
    add_section_title(doc, "09", "启用前检查", "完成录入不等于可以上架；必须先在测试环境把数据、素材和工作台效果一起验收。")
    checklist = [
        "目录路径正确，没有把珠子放进配饰分类。",
        "图片已上传 COS，透明背景、居中与比例正常。",
        "品种形制与实物一致，安装方式正确。",
        "非圆珠的穿线占位、外观宽度和高度已填写。",
        "价格为真实销售价格，不是测试价格。",
        "库存不大于实际可售数量。",
        "工作台、订单预览和材料列表显示正常。",
        "测试环境验证完成后，才按发布流程同步正式环境。",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=5, line_spacing=1.22)
        set_run_font(p.add_run("□  "), size=12, color=GOLD, bold=True)
        set_run_font(p.add_run(item), size=10.5, color=INK)

    add_heading(doc, "运营后台入口")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=7, line_spacing=1.2)
    set_run_font(p.add_run("测试环境："), size=10.5, color=INK, bold=True)
    add_hyperlink(p, "打开测试环境运营后台", "https://api.yustream.cn/test-api/admin")
    add_callout(doc, "环境纪律", "新材料先在测试环境创建和验收；未经发布流程，不直接修改正式环境。", tone="caution")

    add_page_break(doc)
    add_section_title(doc, "10", "常见错误速查", "遇到层级或规格拿不准时，先对照本页再保存。")
    add_text_table(
        doc,
        ["错误操作", "正确处理"],
        [
            ["重复创建“珠子”或“配饰”", "直接复用已有材料类型。"],
            ["把“绿幽灵”建成材料分类", "建为“幽灵水晶”下的品种。"],
            ["把“随形”当作珠径规格", "建为配饰分类或形制，并填写实测尺寸。"],
            ["在品种名称里写价格或库存", "价格和库存只维护在 SKU。"],
            ["随形材料只填最大尺寸", "同时填写穿线占位、外观宽度和高度。"],
            ["图片直接使用临时抠图", "按素材管线处理后上传 COS。"],
            ["资料未完成就启用", "先保持停用，验收工作台效果后再启用。"],
        ],
        [3900, 5460],
    )
    add_callout(doc, "一句话复核", "先确认它属于珠子还是配饰，再判断分类与品种，最后为每个可售单位建立 SKU。")

    set_headers_and_footers(doc)
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
