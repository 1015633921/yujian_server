from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "operations" / "宇涧运营后台使用手册.md"
OUTPUT = ROOT / "docs" / "operations" / "宇涧运营后台使用手册.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BRAND_GREEN = "66745F"
BRAND_GREEN_DARK = "435043"
BRAND_PALE = "EEF2EB"
BRAND_WARM = "F6F2E9"
TABLE_HEADER = "E8EEF5"
TEXT = "232A25"
MUTED = "68716B"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def choose_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    count = len(headers)
    if count == 2:
        return [2700, 6660]
    if count == 3:
        return [2200, 3480, 3680]
    if count == 4:
        return [1700, 2400, 3060, 2200]
    weights = []
    for index, header in enumerate(headers):
        max_len = max([len(header)] + [len(row[index]) for row in rows if index < len(row)])
        weights.append(max(5, min(max_len, 24)))
    total = sum(weights)
    widths = [round(CONTENT_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def set_run_font(run, *, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None, ascii_font: str = "Calibri", east_asia_font: str = "Hiragino Sans GB") -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, color=TEXT)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, color=TEXT, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, color=BRAND_GREEN_DARK, ascii_font="Menlo", east_asia_font="Hiragino Sans GB")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, color=TEXT)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_style_font(style, size: float, color: str, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11, TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BRAND_GREEN, 18, 10),
        "Heading 2": (13, BRAND_GREEN, 14, 7),
        "Heading 3": (12, BRAND_GREEN_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style_font(style, 11, TEXT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("宇涧运营后台使用手册  |  V1.0")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("内部运营参考  ·  第 ")
    set_run_font(r, size=9, color=MUTED)
    add_page_field(p)
    r = p.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("运营标准作业手册")
    set_run_font(r, size=11, color=BRAND_GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("宇涧运营后台使用手册")
    set_run_font(r, size=30, color=BRAND_GREEN_DARK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("从登录、材料与内容维护，到订单履约、售后与仓库管理")
    set_run_font(r, size=14, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    values = [
        ("版本", "V1.0"),
        ("适用版本", "宇涧运营后台 · 2026 年 7 月"),
        ("适用对象", "新运营、客服、仓库、内容编辑、管理员"),
        ("核对日期", "2026 年 7 月 23 日"),
    ]
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], BRAND_PALE)
        for index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, color=BRAND_GREEN_DARK if index == 0 else TEXT, bold=index == 0)
    configure_table_geometry(table, [2700, 6660])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("本地演示数据截图 · 不含真实用户隐私")
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("目录")
    set_run_font(r, size=22, color=BRAND_GREEN_DARK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("按业务流程编排，可从日常操作清单或对应模块开始阅读")
    set_run_font(r, size=10.5, color=MUTED)
    headings = [line[3:].strip() for line in lines if line.startswith("## ")]
    for heading in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        number, _, title = heading.partition(" ")
        r = p.add_run(number)
        set_run_font(r, size=9.8, color=BRAND_GREEN, bold=True)
        r = p.add_run(f"  {title}")
        set_run_font(r, size=9.8, color=TEXT)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), BRAND_WARM)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BRAND_GREEN)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run("提示  ")
    set_run_font(r, size=10.5, color=BRAND_GREEN_DARK, bold=True)
    add_inline_runs(p, text)


def add_image(doc: Document, alt: str, path_text: str) -> None:
    image_path = (SOURCE.parent / path_text).resolve()
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(image_path), width=Inches(6.5))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt)
    doc_pr.set("title", alt)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_together = True
    r = caption.add_run(alt)
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        set_cell_shading(cell, TABLE_HEADER)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=BRAND_GREEN_DARK, bold=True)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_inline_runs(p, value)
            for run in p.runs:
                if run.font.size is None:
                    set_run_font(run, size=9.4, color=TEXT)
    configure_table_geometry(table, choose_widths(headers, rows))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    def cells(line: str) -> list[str]:
        return [part.strip() for part in line.strip().strip("|").split("|")]

    headers = cells(lines[start])
    rows: list[list[str]] = []
    index = start + 2
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        row = cells(lines[index])
        if len(row) < len(headers):
            row.extend([""] * (len(headers) - len(row)))
        rows.append(row[: len(headers)])
        index += 1
    return headers, rows, index


def build() -> Path:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_toc(doc, lines)

    index = 1
    while index < len(lines) and not lines[index].startswith(">"):
        index += 1
    if index < len(lines) and lines[index].startswith(">"):
        add_callout(doc, lines[index][1:].strip())
        index += 1

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        image_match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            index += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:], style="Heading 2")
            index += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(line[3:], style="Heading 1")
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("> "):
            add_callout(doc, line[2:])
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|", lines[index + 1]):
            headers, rows, next_index = parse_table(lines, index)
            add_table(doc, headers, rows)
            index = next_index
            continue
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, text)
            index += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
            index += 1
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, line.rstrip("  "))
        index += 1

    doc.core_properties.title = "宇涧运营后台使用手册"
    doc.core_properties.subject = "面向新运营人员的后台字段说明与操作指南"
    doc.core_properties.author = "宇涧运营团队"
    doc.core_properties.keywords = "宇涧, 运营后台, 使用手册, SOP"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
